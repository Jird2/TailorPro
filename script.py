import openai, os, fitz, pdf2docx, tempfile
from docx import Document
from docx2pdf import convert
import re

"""Suggestions: 
- Incorporate reviews on the front page 
- Bug report section (contact something@gmail.com) ✅
- Donate, cause why not?! :| ✅
- If they want a specific request for resume tailoring -> Paid version -> TailorPro? ($5)
- Play around with the font and color of "ai" letters in "Tailor" in the title to further aesthetic
- Import navbar from Bootstrap (maybe) ✅
- MAKE PDF WORK!!!!
"""

SUMMARY_HEADINGS = [
    "professional summary",
    "summary",
    "intro",
    "career summary",
    "about me",
    "executive summary",
    "profile summary",
    "communication",
    "introduction"
]

EXPERIENCE_HEADINGS = [
    "experience",
    "work experience",
    "professional experience",
    "employment history",
    "career history",
    "work history"
]

SKILLS_HEADINGS = [
    "skills",
    "technical skills",
    "core competencies",
    "key skills",
    "expertise",
    "proficiencies"
]

def tailor_summary_to_resume(current_summary, company_description, job_posting):
    prompt = f"""
You are a professional resume editor.
Here is a user's current resume summary:
{current_summary}

Here is a company description:
{company_description}

Here is a job posting:
{job_posting}

Rewrite the resume summary to align better with the company's mission, values, and the job responsibilities.
Use relevant keywords and keep it professional and concise (3-5 sentences).
Avoid mentioning the company name directly. Subtly tailor the language to match the company's goals.
"""
    client = openai.OpenAI()
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response.choices[0].message.content.strip()

def summary_from_docx(file):
    doc = Document(file)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() in SUMMARY_HEADINGS:
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    return doc.paragraphs[j].text.strip()
    for p in doc.paragraphs:
        if p.text.strip():
            return p.text.strip()
    return ""

def update_summary_in_docx(file, new_summary, output):
    doc = Document(file)
    found_summary_heading = False
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() in SUMMARY_HEADINGS:
            found_summary_heading = True
            continue
        if found_summary_heading and paragraph.text.strip():
            if paragraph.runs:
                font_name = paragraph.runs[0].font.name
                font_size = paragraph.runs[0].font.size
            else:
                font_name = None
                font_size = None
            paragraph.clear()
            run = paragraph.add_run(new_summary)
            if font_name: run.font.name = font_name
            if font_size: run.font.size = font_size
            doc.save(output)
            return
    paragraph = doc.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(new_summary)
    doc.save(output)

def extract_resume_content(doc_path):
    """Extract all content from resume for comprehensive analysis"""
    doc = Document(doc_path)
    content = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content += paragraph.text.strip() + "\n"
    return content

def generate_resume_suggestions(doc_path, company_desc, job_posting):
    """Generate suggestions based on the resume content"""
    resume_content = extract_resume_content(doc_path)
    
    client = openai.OpenAI()
    prompt = f"""
    You are a resume consultant analyzing a resume for improvement opportunities.

    Here is the complete resume content:
    {resume_content}

    Company Description:
    {company_desc}

    Job Posting:
    {job_posting}

    Analyze the resume and provide specific suggestions for:
    1. Skills that should be highlighted or added (based on what's already there)
    2. Ways to strengthen existing bullet points and descriptions
    3. Keywords to incorporate
    4. Any structural improvements

    Focus on enhancing what already exists rather than inventing new experiences.
    Be specific about which sections need improvement and how.
    """
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response.choices[0].message.content.strip()

def is_bullet_point(text):
    """Improved bullet point detection"""
    text = text.strip()
    
    # Return False for very short text or empty lines
    if len(text) < 10:
        return False
    
    # Common bullet symbols and patterns
    bullet_patterns = [
        r'^[-•·▪▫◦‣⁃★▸►→]',
        r'^\d+\.',
        r'^[a-zA-Z]\.',
        r'^\([a-zA-Z0-9]+\)',
    ]
    
    for pattern in bullet_patterns:
        if re.match(pattern, text):
            return True
    
    # Check for action-oriented sentences (common in experience sections)
    action_words = [
        'developed', 'managed', 'led', 'created', 'implemented', 'designed', 
        'analyzed', 'coordinated', 'supervised', 'executed', 'established',
        'improved', 'optimized', 'delivered', 'achieved', 'collaborated',
        'maintained', 'operated', 'assisted', 'supported', 'facilitated',
        'organized', 'planned', 'trained', 'conducted', 'performed'
    ]
    
    # If it starts with an action word and is a reasonable length, likely a bullet
    first_word = text.split()[0].lower() if text.split() else ""
    if first_word in action_words and 20 <= len(text) <= 300:
        return True
    
    return False

def is_section_header(text):
    """Improved section header detection"""
    text_clean = text.lower().strip()
    
    # Remove common formatting characters
    text_clean = re.sub(r'[:\-_=]+$', '', text_clean)
    
    all_headers = (
        SUMMARY_HEADINGS + 
        EXPERIENCE_HEADINGS + 
        SKILLS_HEADINGS + 
        ['education', 'certifications', 'projects', 'achievements', 'awards',
         'publications', 'volunteer', 'languages', 'interests', 'references']
    )
    
    # Exact matches
    if text_clean in all_headers:
        return True
    
    # Partial matches for longer headers
    for header in all_headers:
        if header in text_clean and len(text_clean) <= len(header) + 10:
            return True
    
    # Check if its a short line (prob a header) thats not a bullet point
    if len(text.strip()) <= 50 and not is_bullet_point(text) and text.strip().isupper():
        return True
    
    return False

def preserve_formatting(original_paragraph):
    """Extract formatting from original paragraph"""
    formatting = {}
    if original_paragraph.runs:
        first_run = original_paragraph.runs[0]
        formatting = {
            'font_name': first_run.font.name,
            'font_size': first_run.font.size,
            'bold': first_run.font.bold,
            'italic': first_run.font.italic,
            'underline': first_run.font.underline
        }
    return formatting

def apply_formatting(run, formatting):
    """Apply formatting to a run"""
    try:
        if formatting.get('font_name'): 
            run.font.name = formatting['font_name']
        if formatting.get('font_size'): 
            run.font.size = formatting['font_size']
        if formatting.get('bold'): 
            run.font.bold = formatting['bold']
        if formatting.get('italic'): 
            run.font.italic = formatting['italic']
        if formatting.get('underline'): 
            run.font.underline = formatting['underline']
    except Exception as e:
        print(f"Warning: Could not apply formatting: {e}")

def revise_summary_and_bullets(doc_path, suggestions, job_posting, output_path):
    """
    Comprehensively enhance the resume based on AI suggestions
    """
    client = openai.OpenAI()
    
    # Create a fresh copy of the document
    doc = Document(doc_path)
    
    # Track sections and content
    current_section = None
    summary_processed = False
    modifications_made = 0
    
    print("Starting comprehensive resume enhancement...")
    
    # Process each paragraph
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        
        if not para_text:
            continue
        
        # Skip if this looks like inserted junk
        if ("job requirements:" in para_text.lower() or 
            "position prerequisites:" in para_text.lower() or
            para_text.lower().startswith("original:")):
            para.clear()
            continue
        
        # Identify section headers
        if is_section_header(para_text):
            current_section = para_text.lower()
            print(f"[{i}] Section: {current_section}")
            continue
        
        # Process content based on type
        try:
            enhanced_text = None
            modification_type = None
            
            # Handle summary/intro sections
            if (any(summary_word in (current_section or "") for summary_word in SUMMARY_HEADINGS) 
                and not summary_processed and len(para_text) > 20):
                
                print(f"[{i}] Processing summary: {para_text[:50]}...")
                
                prompt = f"""
Rewrite this professional summary to better align with this job opportunity:

Current Summary: {para_text}

Job Requirements: {job_posting}

Rules:
- Keep it concise (3-4 sentences max)
- Include relevant keywords naturally
- Maintain the person's experience level and achievements
- Make it more relevant to the target role
- Don't mention specific company names

Return only the enhanced summary:"""

                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.6
                )
                enhanced_text = response.choices[0].message.content.strip()
                modification_type = "summary"
                summary_processed = True
            
            # Handle bullet points in experience sections
            elif is_bullet_point(para_text):
                
                print(f"[{i}] Processing bullet point: {para_text[:50]}...")
                
                prompt = f"""
Enhance this resume bullet point to be more impactful for the target job:

Original: {para_text}

Target Job: {job_posting}

Rules:
- Start with a strong action verb
- Include quantifiable results when possible (but don't invent numbers)
- Use keywords from the job posting naturally
- Keep the same general structure and truthfulness
- Make it more compelling and specific
- Maintain any existing bullet formatting

Return only the enhanced bullet point:"""

                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.6
                )
                enhanced_text = response.choices[0].message.content.strip()
                modification_type = "bullet"
            
            # Apply enhancement if made
            if enhanced_text:
                """Clean up AI response"""
                enhanced_text = re.sub(r'^(Enhanced\s+)?(Summary|Bullet\s+Point):?\s*', '', enhanced_text, flags=re.IGNORECASE)
                enhanced_text = enhanced_text.strip().strip('"\'')
                
                """Original formatting"""
                formatting = preserve_formatting(para)
                
                """Replace content"""
                para.clear()
                new_run = para.add_run(enhanced_text)
                apply_formatting(new_run, formatting)
                
                modifications_made += 1
                print(f"Enhanced {modification_type} #{modifications_made}")
        
        except Exception as e:
            print(f"Error enhancing paragraph {i}: {e}")
            continue
    
    """Save the enhanced document"""
    doc.save(output_path)
    print(f"Enhanced resume saved to {output_path}")
    print(f"Total modifications made: {modifications_made}")

def run(resume_path, company_description, job_posting, output):
    ext = resume_path.lower().split('.')[-1]
    print("File extension detected:", ext)
    if ext == "docx":
        current_summary = summary_from_docx(resume_path)
        updated_summary = tailor_summary_to_resume(current_summary, company_description, job_posting)
        update_summary_in_docx(resume_path, updated_summary, output)
    else:
        raise ValueError("Unsupported file type. Use .docx only.")
    print("Resume tailored and saved to", output)